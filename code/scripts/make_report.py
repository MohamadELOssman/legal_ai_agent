#!/usr/bin/env python3
"""Generate a Word progress report for the thesis-advisor meeting."""

from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("../Legal_AI_Project_Update.docx")  # repo root: legal_ai_agent/

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def para(text, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Lebanese Legal AI — Progress Report")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A Multi-Agent, Retrieval-Augmented System for Lebanese Penal Law")
r.font.size = Pt(12)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Prepared for the thesis-advisor meeting · {date.today():%B %d, %Y}").italic = True

doc.add_paragraph()

# ── 1. Overview ──────────────────────────────────────────────────────────────
h("1. Overview", 1)
para("This report summarises the work completed on the Lebanese Legal AI system. "
     "The system answers legal questions in Arabic, English, and French by retrieving "
     "the relevant articles of the Lebanese Penal Code and court rulings, then writing "
     "a structured legal memorandum through a pipeline of specialised AI agents.")
para("The focus of this phase was two-fold: (1) making the system more accurate, "
     "trustworthy, and reliable, and (2) building a rigorous benchmark so that every "
     "design choice is backed by evidence rather than assumption. The benchmark is "
     "described in detail in Section 3, as it is the core scientific contribution.")

# ── 2. What's new ────────────────────────────────────────────────────────────
h("2. What Was Added", 1)

h("2.1 Models", 2)
bullet(" Standardised the whole system on Claude Sonnet 4.5, and added a model "
       "selector so each run can use Sonnet 4.5, Sonnet 4.6, Opus 4.6, or Haiku 4.5 "
       "(useful for comparing models in the thesis).", "Up-to-date models. ")

h("2.2 Data and Corpus", 2)
bullet(" Added the English Penal Code (242 articles) alongside the Arabic one, making "
       "the searchable corpus bilingual (713 documents: 417 Arabic + 242 English articles "
       "+ 54 court rulings).", "Bigger, bilingual corpus. ")
bullet(" Built a reproducible pipeline that reads the source documents, indexes them, and "
       "records exactly what was indexed — so the corpus can be rebuilt with one command "
       "and new sources (e.g., contract law, French texts) can be added automatically.",
       "Reproducible ingestion. ")

h("2.3 Accuracy and Trust", 2)
bullet(" Each legal point the system makes is now checked against the source documents; "
       "anything not found in the corpus is flagged, and a 'hallucination rate' is reported "
       "for every answer.", "Grounding / anti-hallucination. ")
bullet(" Every citation (e.g., 'Article 549') is verified against a master list of real "
       "article numbers, so the answer shows which citations are confirmed.",
       "Citation verification. ")
bullet(" Fixed several internal bugs where agents were silently working on empty input, "
       "which previously weakened the legal reasoning.", "Correctness fixes. ")

h("2.4 User Experience", 2)
bullet(" The memorandum is now written entirely in the language of the question "
       "(Arabic in, Arabic out), and Arabic documents display right-to-left in a clean, "
       "professional layout.", "Language-matched, RTL output. ")
bullet(" Each answer shows its sources, a cost/time breakdown, and trust indicators.",
       "Transparency. ")

h("2.5 Engineering", 2)
bullet(" Added an automated test suite and continuous integration so changes are checked "
       "automatically, plus per-answer cost, speed, and token tracking.",
       "Tests, CI, and monitoring. ")

# ── 3. Benchmarking (main section) ───────────────────────────────────────────
h("3. Benchmarking — The Core of This Phase", 1)

h("3.1 Why a benchmark was needed", 2)
para("A legal AI system is only credible if its accuracy can be measured. Without a "
     "benchmark, we can only say the answers 'look good'. With one, we can prove how often "
     "the system finds the correct law, how accurate its citations are, and whether the "
     "multi-agent design actually outperforms a simple single-AI baseline. The benchmark "
     "turns subjective impressions into objective, repeatable numbers — and lets us test "
     "every design decision with evidence.")

h("3.2 How the benchmark dataset was built", 2)
para("I generated a dataset of 196 legal questions, and crucially, every question is "
     "'grounded' in the real corpus:")
bullet(" Questions are generated directly from the actual Penal Code articles and court "
       "rulings — not invented — so they reflect real legal content.", "Grounded in real law. ")
bullet(" Each question is stored with the correct article number(s) it should rely on (the "
       "'gold answer'), and every gold label was automatically validated against the corpus.",
       "Verified gold answers. ")
bullet(" The set is trilingual (≈69 Arabic, 65 English, 62 French) to test the system's "
       "multilingual ability.", "Three languages. ")
bullet(" It mixes two question types: general legal questions ('what does the law say?') "
       "and case scenarios ('here are the facts — assess them').", "Two realistic question types. ")
bullet(" It covers 204 different articles, so the test is broad, not narrow.", "Broad coverage. ")

h("3.3 What the benchmark measures", 2)
para("The system is evaluated on two layers — first whether it finds the right law, then "
     "whether it answers well:")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for c, t in zip(hdr, ["Layer", "Metrics", "What it tells us"]):
    c.paragraphs[0].add_run(t).bold = True
rows = [
    ("Retrieval\n(finding the law)",
     "Precision@k, Recall@k, MRR, nDCG@k",
     "How often, and how highly ranked, the correct articles appear in the results."),
    ("Citations\n(accuracy)",
     "Precision, Recall, F1 vs. gold articles",
     "Whether the answer cites the correct articles — measured automatically."),
    ("Answer quality",
     "LLM-as-judge score 1–5 (legal correctness, citation quality, completeness, clarity)",
     "Whether the final memorandum is legally sound and well written."),
    ("Efficiency",
     "Latency, tokens, cost per query",
     "Whether the system is practical to run."),
    ("Statistical rigour",
     "Mean ± 95% confidence interval; paired significance tests",
     "Whether differences between methods are real, not chance."),
]
for a, b, c in rows:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(a).bold = True
    cells[1].text = b
    cells[2].text = c

doc.add_paragraph()

h("3.4 Why I designed it this way", 2)
bullet(" Because legal answers must point to specific articles, the gold answer is the set "
       "of correct article numbers. This makes scoring objective and automatic — no human "
       "needed for the core retrieval and citation metrics.", "Article-level gold answers. ")
bullet(" Grounding the questions in the real corpus guarantees every question has a known, "
       "checkable correct answer — avoiding made-up questions with no verifiable answer.",
       "Grounded generation. ")
bullet(" Three languages and two question types reflect how the system will actually be "
       "used, so the score represents real-world performance.", "Realistic coverage. ")
bullet(" Comparing against single-AI baselines (with and without retrieval) directly tests "
       "the thesis claim that a multi-agent design is better.", "Baseline comparison. ")
bullet(" Reporting confidence intervals and significance tests ensures conclusions are "
       "scientifically defensible, not based on a single lucky run.", "Statistical honesty. ")

h("3.5 What the benchmark already revealed", 2)
para("Running the benchmark immediately produced evidence-based decisions:")
bullet(" Hybrid search (keyword + meaning) retrieves the correct law more often than "
       "meaning-only search — so it is now the default.", "Hybrid retrieval wins. ")
bullet(" A popular English 're-ranking' model actually made results worse on this "
       "Arabic/legal corpus, so it was disabled. The benchmark caught this; intuition "
       "would have kept it.", "A common technique was rejected. ")
bullet(" The local embedding model already in use outperformed two larger alternatives — "
       "validating the existing choice with data.", "Embedding choice validated. ")

# ── 4. Results snapshot ──────────────────────────────────────────────────────
h("4. Current Results (Snapshot)", 1)
para("Numbers below are from the 196-question benchmark. Retrieval is fully measured; "
     "answer-quality is an early sample and will be finalised in a full run.", italic=True)

t2 = doc.add_table(rows=1, cols=3)
t2.style = "Light Grid Accent 1"
for c, t in zip(t2.rows[0].cells, ["Aspect", "Result", "Reading"]):
    c.paragraphs[0].add_run(t).bold = True
res = [
    ("Finds correct law (Recall@5)", "~50%", "The correct article is in the top 5 about half the time — solid, with clear room to improve."),
    ("Answer quality (judge, 1–5)", "≈4.6 / 5 (early sample)", "The written memoranda are rated as high quality."),
    ("Citation precision", "Low (being improved)", "It cites the right legal area but not always the exact article — the next target."),
    ("Speed / cost", "~3 min, ~$0.12 per full answer", "Practical for research use."),
]
for a, b, c in res:
    cells = t2.add_row().cells
    cells[0].paragraphs[0].add_run(a).bold = True
    cells[1].text = b
    cells[2].text = c

doc.add_paragraph()
para("Interpretation: the system writes strong legal memoranda, and its main limitation "
     "is pinpoint citation accuracy, which is bounded by how often the correct article is "
     "retrieved. This directly sets the priorities below.")

# ── 5. Next steps ────────────────────────────────────────────────────────────
h("5. Next Steps", 1)
bullet(" Improve retrieval so the correct article is found more often (tune the search; "
       "test better multilingual models).", "1. Raise retrieval recall. ")
bullet(" Make the system cite only the directly-applicable articles, not neighbouring ones.",
       "2. Improve citation precision. ")
bullet(" Run the full benchmark to produce final, statistically-backed comparison numbers "
       "(multi-agent vs. baselines).", "3. Full evaluation run. ")
bullet(" Expand the corpus to contract law and French legal texts (the pipeline already "
       "supports adding them).", "4. Broaden the corpus. ")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.add_run("All work is version-controlled, tested, and reproducible.").italic = True

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Saved: {OUT.resolve()}")
print(f"Paragraphs: {len(doc.paragraphs)} | Tables: {len(doc.tables)}")
