#!/usr/bin/env python3
"""Detailed Word reference document for the Lebanese Legal AI project (for the professor)."""

from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGS = Path("experiments/figures")
OUT = Path("../Legal_AI_Full_Reference.docx")

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GREEN = RGBColor(0x0F, 0x7A, 0x3E)
GREY = RGBColor(0x33, 0x33, 0x33)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)


def H(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY
    return p


def para(text="", italic=False, size=11, bold=False, space=6):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space)
    return p


def bullet(text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        p.add_run(lead).bold = True
    p.add_run(text)
    return p


def code_block(lines):
    """A light-shaded monospace block for prompt excerpts."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.2); p.paragraph_format.space_after = Pt(8)
    for i, ln in enumerate(lines):
        r = p.add_run(("" if i == 0 else "\n") + ln)
        r.font.name = "Consolas"; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for r in cells[j].paragraphs[0].runs:
                r.font.size = Pt(10)
                if j == 0:
                    r.bold = True
    doc.add_paragraph()
    return t


# ── Title page ────────────────────────────────────────────────────────────────
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Lebanese Legal AI — Detailed Project Reference"); r.bold = True
r.font.size = Pt(22); r.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("A Multi-Agent, Retrieval-Augmented System for Lebanese Legal Research")
r.font.size = Pt(13); r.font.color.rgb = GREY
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
d.add_run(f"Reference document for the advisor presentation · {date.today():%B %d, %Y}").italic = True
para("", space=4)
para("This document is a complete, plain-language reference for the project. It mirrors the "
     "presentation and adds the underlying detail — the agents and their prompts, the retrieval "
     "method, the trust mechanisms, the web application, and especially the benchmarking process "
     "(question generation, the evaluator/judge agent, and the source-of-truth answer). Use it to "
     "answer any question during the meeting.", italic=True)
para("The system is delivered two ways over the same specialised agents: (a) a conversational "
     "CHAT ASSISTANT driven by an agentic orchestrator that calls only the sub-agents a question "
     "actually needs, and (b) a fixed END-TO-END PIPELINE that always runs every agent for full "
     "transparency. It covers two Lebanese codes — the Penal Code and the Code of Criminal "
     "Procedure (قانون أصول المحاكمات الجزائية) — and defaults to Claude Sonnet 5.", italic=True)

# ── 1. Overview ────────────────────────────────────────────────────────────────
H("1. Overview & Motivation", 1)
para("Lebanese law is trilingual: Arabic is the official language, while French and English are "
     "widely used in practice. Legal sources are scattered and difficult to search, ordinary "
     "citizens struggle to understand their rights, and lawyers spend hours locating the correct "
     "articles and precedents. General-purpose chatbots are unreliable for legal work because they "
     "invent article numbers (hallucinate).")
para("This project answers legal questions by first RETRIEVING the correct law from a curated "
     "corpus, then writing a grounded, cited answer whose accuracy can be measured. It is delivered "
     "as a web application (a chat assistant plus a full pipeline view) with a rigorous, "
     "reproducible benchmark. The corpus covers the Lebanese Penal Code and the Code of Criminal "
     "Procedure, in Arabic (primary) and English.")

H("Research objectives", 2)
bullet("can specialised agents outperform a single large language model?", "Multi-agent design: ")
bullet("how well can the system find the right law across Arabic / English / French?", "Trilingual retrieval: ")
bullet("can every legal claim and citation be grounded and verified against the corpus?", "Trustworthiness: ")
bullet("can the output adapt to the user — citizen, lawyer, or judge?", "Adaptivity: ")
bullet("build a benchmark that proves accuracy objectively, against a source of truth.", "Rigorous evaluation: ")

# ── 2. Architecture ────────────────────────────────────────────────────────────
H("2. System Architecture — The Agent Pipeline", 1)
para("The system is a pipeline of specialised agents. Each agent has one job and passes its output "
     "to the next. A question flows through the following steps:")
table(["#", "Agent", "Role"], [
    ("0", "Orchestrator", "Classifies the question type AND the user type; routes the whole pipeline."),
    ("1", "Query Understanding", "Detects language, domain, and facts as validated structured data."),
    ("2", "Research (RAG)", "Retrieves the most relevant articles and rulings using hybrid search."),
    ("3", "Analysis", "Extracts the applicable provisions, grounded against the retrieved text."),
    ("4", "Reasoning", "Applies the law to the facts (chain-of-thought)."),
    ("5", "Citation", "Formats and verifies every article citation against the corpus."),
    ("6", "Writing", "Writes the final answer in the user's shape and language."),
])
para("A separate trust layer runs alongside the pipeline: it flags any legal claim not found in "
     "the sources and reports a hallucination rate for each answer.")

H("2.1 Two execution modes: fixed pipeline vs. agentic chat", 2)
para("The same agents are used in two ways:")
bullet("always runs all seven agents in order. Maximum transparency — every step, its input/output, "
       "the retrieved documents, and per-agent cost/latency are shown. Best for demonstration and audit.",
       "Fixed pipeline: ")
bullet("a conversational assistant where the Orchestrator is given the sub-agents as TOOLS and calls "
       "ONLY the ones a given question needs — a simple citizen question may need one Research call; a "
       "complex case may chain Research → Analysis → Citation. This saves time and tokens and behaves "
       "more intelligently, while the named sub-agents (Research, Analysis, Citation) still run and are "
       "shown live, step by step. Multi-turn, with saved conversations.", "Agentic chat: ")
para("Sub-agents exposed to the chat orchestrator (as callable tools):", bold=True, space=2)
code_block([
    "research_agent(query)          -> relevant articles (both codes) + court rulings",
    "analysis_agent(question)       -> applicable provisions, grounded and explained",
    "citation_agent(article_nums)   -> verifies article numbers against the corpus",
])
para("Every retrieved item is labelled with which code it comes from, because the two codes share "
     "article numbers (both have an Article 24, 90, 233, …).")

# ── 3. System prompts ──────────────────────────────────────────────────────────
H("3. How the Agents Are Instructed (System Prompts)", 1)
para("Each agent is controlled by a role-specific system prompt that fixes its behaviour. Two design "
     "choices make the system reliable:")
bullet("the entry agents return VALIDATED structured objects via tool-use (a fixed schema), instead "
       "of free text — so there is no fragile parsing of the model output.", "Structured outputs: ")
bullet("every generative prompt forbids inventing law and locks the output language.", "Grounding rules: ")

para("Orchestrator (system prompt, excerpt):", bold=True, space=2)
code_block([
    "You are the Orchestrator of a Lebanese Legal AI system.",
    "Read the user input, classify it, and return a routing decision.",
    "You never answer legal questions directly. Classify TWO things:",
    "  A) query_type : general_legal_query | case_analysis",
    "  B) user_type  : citizen | lawyer | judge   (who is asking)",
])
para("Writing / Analysis agents (rules, excerpt):", bold=True, space=2)
code_block([
    "Only cite articles that are supplied to you; never invent legal references.",
    "Write the ENTIRE memorandum — including headings — in the query's language ONLY.",
])
para("Routing schema (enforced via tool-use):", bold=True, space=2)
code_block([
    "class RoutingDecision:",
    "    query_type: str            # general_legal_query | case_analysis",
    "    user_type:  str            # citizen | lawyer | judge",
    "    detected_language: str     # ar | en | fr",
    "    legal_domain: str          # criminal | civil | ...",
    "    extracted_facts: list[str]",
    "    pipeline_config: {...}      # research / analysis / writing modes",
])

# ── 4. Adaptive output ─────────────────────────────────────────────────────────
H("4. Adaptive Output — Who Is Asking?", 1)
para("The Orchestrator detects the user type (or the user selects it) and the Writing agent produces "
     "the matching output shape. Each shape uses a FIXED set of section headers, written verbatim and "
     "in order (validated by a domain expert). The Arabic headers below are authoritative; French and "
     "English use faithful equivalents so the structure is identical across languages.")

para("Citizen — a plain question (two parts):", bold=True, space=2)
code_block(["التحليل القانوني        (Legal Analysis)",
            "الإجابة باختصار         (Answer in Brief)"])

para("Lawyer — defending a client's case:", bold=True, space=2)
code_block(["الوقائع المنتجة              (Relevant Facts)",
            "القوانين والمواد ذات الصلة   (Applicable Laws & Articles)",
            "طريقة الدفاع                 (Defense Strategy)",
            "الإجابة باختصار             (Answer in Brief)"])

para("Judge — facts given, ruling expected:", bold=True, space=2)
code_block(["المحكمة المختصة              (Competent Court)",
            "أطراف الدعوى                (Parties to the Case)",
            "الوقائع المنتجة              (Relevant Facts)",
            "القوانين والمواد ذات الصلة   (Applicable Laws & Articles)",
            "تطبيق القانون على الوقائع    (Application of Law to Facts)",
            "الحكم                        (Judgment)"])

para("Case study — a general question / neutral analysis (no specific role):", bold=True, space=2)
code_block(["الوقائع المنتجة              (Relevant Facts)",
            "الاشكالية القانونية          (Legal Issue)",
            "المواد والقوانين ذات الصلة   (Applicable Laws & Articles)",
            "تطبيق القانون على الوقائع    (Application of Law to Facts)",
            "الحل                         (Solution)",
            "المحكمة المختصة (في حال وجودها)  (Competent Court, if applicable)"])

para("Auto-detect infers the role from the phrasing (e.g. 'my client' → lawyer, 'render the verdict' "
     "→ judge). Both the fixed pipeline and the agentic chat enforce these same templates.")

# ── 5. Corpus ──────────────────────────────────────────────────────────────────
H("5. Corpus & Data Foundation", 1)
bullet("two Lebanese codes plus case law — Penal Code (417 Arabic + 242 English articles), Code of "
       "Criminal Procedure (431 Arabic articles), and 54 Court of Cassation rulings = 1,144 indexed "
       "documents.", "Corpus: ")
table(["Source", "Articles / docs", "Language"], [
    ("Penal Code (قانون العقوبات)", "659 (417 AR + 242 EN)", "Arabic + English"),
    ("Code of Criminal Procedure (أصول المحاكمات الجزائية)", "431", "Arabic"),
    ("Court of Cassation rulings", "54", "Arabic"),
    ("Total indexed", "1,144 documents", "—"),
])
bullet("a reproducible pipeline (build_index.py) AUTO-DISCOVERS every code file in the documents "
       "folder, embeds them into a searchable index, regenerates the citation article index per code, "
       "and records exactly what was indexed (a manifest).", "Ingestion: ")
bullet("the Code of Criminal Procedure was added simply by dropping its JSON in the corpus folder and "
       "re-running the pipeline — no code change was needed for ingestion. Contract law / French texts "
       "can be added the same way.", "Extensible: ")
bullet("each code is tagged with its document type; because the two codes share article numbers, every "
       "answer states WHICH code a cited article belongs to.", "Disambiguation: ")
bullet("every article keeps a validated number, later used as the 'gold answer' in the benchmark.",
       "Provenance: ")

# ── 6. Retrieval ───────────────────────────────────────────────────────────────
H("6. Retrieval (RAG)", 1)
bullet("Hybrid search combines keyword matching (BM25) with meaning-based matching (dense "
       "embeddings). This was chosen because the benchmark showed it wins.", "Method: ")
bullet("a local, free multilingual sentence-transformer (mpnet).", "Embeddings: ")
bullet("articles and rulings are searched in separate pools; matching by article number makes "
       "results language-agnostic (an English query can correctly match the Arabic article).",
       "Cross-lingual: ")
bullet("a popular English re-ranking model was tested and REJECTED because it made results worse on "
       "this Arabic/legal corpus — a decision made by measurement, not intuition.", "Evidence-based: ")

# ── 7. Trust ───────────────────────────────────────────────────────────────────
H("7. Trust & Anti-Hallucination", 1)
bullet("each extracted provision is checked against the retrieved text; ungrounded ones are flagged.",
       "Grounding: ")
bullet("every cited article number is verified against a master corpus index — no invented citations.",
       "Citation verification: ")
bullet("the Writing agent receives a CLOSED set and may cite ONLY those verified article numbers.",
       "Closed citation set: ")
bullet("each answer reports a grounding rate and a hallucination rate.", "Trust metric: ")
code_block([
    "CITATION CONSTRAINT: You may cite ONLY these article numbers: 547, 549.",
    "Do NOT mention or cite any other article number, even if related.",
])

# ── 8. Web app ─────────────────────────────────────────────────────────────────
H("8. The Web Application (UI)", 1)
para("Built with Streamlit; it has four areas, selected from a collapsible sidebar:")
bullet("the primary interface — a professional, multi-turn chat. The agentic orchestrator calls only "
       "the sub-agents it needs and shows them running LIVE, step by step (Research / Analysis / "
       "Citation). Answers render Markdown with correct right-to-left layout for Arabic. Each answer "
       "carries a light footer (sub-agent calls · latency · tokens · estimated spend · verified "
       "citations) and a collapsed, professional SOURCES panel — the exact articles/rulings used, each "
       "labelled with its code and marked when actually cited. Conversations are saved to disk (new "
       "chat, rename-by-first-question, delete) and survive restarts.", "Chat Assistant: ")
bullet("ask a question, pick your role, run all 7 agents, and see the memorandum, its sources, trust "
       "indicators, and per-agent cost/latency. The Arabic memo renders right-to-left in a clean "
       "document view.", "End-to-End Pipeline: ")
bullet("run and inspect each agent in isolation (inputs, outputs, retrieved chunks) for debugging.",
       "Individual Agents: ")
bullet("generate test questions, enter reference (source-of-truth) answers, and score the system "
       "(described next).", "Benchmarking: ")

# ── 9. BENCHMARKING ────────────────────────────────────────────────────────────
H("9. Benchmarking (Core Contribution)", 1)
para("A legal AI is only credible if its accuracy can be measured, not just claimed. The benchmark "
     "turns subjective impressions into objective, repeatable numbers, and lets every design "
     "decision be tested with data.")

H("9.1 Generating the benchmark questions", 2)
para("A dedicated generation step builds questions that are GROUNDED in the real corpus, so each "
     "question has a known, checkable correct answer:")
bullet("real Penal-Code articles (and rulings) are sampled from the corpus.", "Step 1 — ")
bullet("the model writes realistic questions answerable from THOSE articles, in a target language.", "Step 2 — ")
bullet("each question is returned WITH the article numbers it is based on — the gold answer.", "Step 3 — ")
bullet("gold numbers are validated against the corpus index; invalid ones are dropped.", "Step 4 — ")
bullet("generation is batched (~10 questions per model call) for efficiency, then deduplicated.", "Step 5 — ")
para("Generation prompt (excerpt):", bold=True, space=2)
code_block([
    "From these Lebanese Penal Code articles:",
    "  - Article 549: ...",
    "  - Article 547: ...",
    "Generate N diverse questions in {language}.",
    "Set gold_articles to those article numbers (choose only from: 547, 548, 549).",
    "Mix about half general questions and half case scenarios.",
])
para("The delivered dataset has 196 questions: ~69 Arabic, 65 English, 62 French; 148 article-based "
     "and 48 ruling-based; covering 204 distinct articles.")

H("9.2 The Evaluator Agent (LLM-as-Judge) and the Source of Truth", 2)
para("Answer quality is scored by a separate, independent LLM (the 'judge'), distinct from the "
     "answering agents. Crucially, the judge does NOT score on its own opinion — it compares each "
     "answer against a human-provided SOURCE-OF-TRUTH (reference) answer.")
bullet("an independent LLM (Claude) at temperature 0 for consistent, deterministic scoring.", "What it is: ")
bullet("the user enters the correct expert answer per question in the benchmark table — the source of truth.",
       "Depends on the reference: ")
bullet("the judge compares the AI answer AGAINST that reference on four dimensions.", "How it scores: ")
bullet("without a reference the judge might score on a biased or wrong opinion; with it, evaluation is "
       "grounded in human expertise — objective and reproducible.", "Why it matters: ")
bullet("the reference answer is REQUIRED for a judged run; the run is blocked until every evaluated "
       "question has one.", "Mandatory: ")
para("Judge prompt (reference-based, excerpt):", bold=True, space=2)
code_block([
    "You are a Lebanese legal evaluation expert.",
    "Compare the AI answer against the REFERENCE (ground-truth) answer.",
    "REFERENCE answer: \"...\"   (entered by the user)",
    "AI answer: \"...\"",
    "Score each 1-5, judged AGAINST the reference:",
    "  legal_correctness, citation_quality, completeness, clarity",
    "Return ONLY JSON.",
])
para("In parallel, an OBJECTIVE citation metric (precision / recall / F1 of the cited articles vs. "
     "the gold article numbers) is computed automatically to cross-check the judge.")

H("9.3 What is measured", 2)
table(["Layer", "Metrics", "What it tells us"], [
    ("Retrieval", "Precision@k, Recall@k, MRR, nDCG, hit-rate", "Does it find the correct article?"),
    ("Citations", "Precision, Recall, F1 vs. gold articles", "Does the answer cite the right articles?"),
    ("Answer quality", "LLM-judge 1–5 (correctness, citations, completeness, clarity)", "Is the answer sound & well written?"),
    ("Reference-based", "Judge vs. the human source-of-truth answer", "Does it match the expert answer?"),
    ("Efficiency", "Latency, tokens, cost per query", "Is it practical to run?"),
    ("Statistics", "Mean ± 95% CI, paired significance tests", "Are differences real, not chance?"),
])
para("Metric definitions in plain terms:", bold=True, space=2)
bullet("fraction of the correct (gold) articles that appear in the top-k retrieved results.", "Recall@k: ")
bullet("did AT LEAST ONE correct article appear in the top-k (a lenient, practical measure).", "Hit-rate@k: ")
bullet("of the articles the answer cites, how many are correct (precision) / of the correct ones, "
       "how many were cited (recall); F1 balances the two.", "Citation P / R / F1: ")

H("9.4 How a benchmark run works in the app", 2)
bullet("set the number of questions; the generator produces grounded questions, shown 10 per page.", "1 · Generate: ")
bullet("review them and type the source-of-truth answer in the editable table (required).", "2 · Reference: ")
bullet("choose systems — Multi-Agent vs. Single-Agent vs. No-RAG.", "3 · Select: ")
bullet("each system answers; the judge scores every answer against the reference.", "4 · Run: ")
bullet("per-system scores, a per-dimension table, charts, citation metrics, and a downloadable JSON.", "5 · Results: ")

# ── 10. Results ────────────────────────────────────────────────────────────────
H("10. Results", 1)
para("The following numbers are from the 196-question benchmark. Retrieval is fully measured; "
     "answer-quality is an early sample.", italic=True)

H("10.1 Retrieval accuracy by language (hit-rate)", 2)
table(["Language", "Hit-rate @5", "Hit-rate @10"], [
    ("Arabic (primary legal language)", "72.5%", "81.2%"),
    ("English", "43.1%", "49.2%"),
    ("French", "38.7%", "46.8%"),
])
para("On Arabic — the official language of Lebanese law — the system surfaces a correct article about "
     "73% of the time within the top 5 (81% within the top 10). English and French are lower because "
     "the corpus is Arabic-primary; closing this cross-lingual gap is the main next step.")
if (FIGS / "fig_languages.png").exists():
    doc.add_picture(str(FIGS / "fig_languages.png"), width=Inches(4.8))

H("10.2 Which search method is best", 2)
table(["Method", "Recall@5", "Precision@5", "MRR", "nDCG@5"], [
    ("Hybrid (BM25 + dense) — default", "0.498", "0.112", "0.360", "0.407"),
    ("Semantic (meaning only)", "0.428", "0.100", "0.324", "0.359"),
    ("Semantic + re-ranker", "0.323", "0.074", "0.187", "0.222"),
    ("Hybrid + re-ranker", "0.327", "0.070", "0.181", "0.219"),
])
para("Hybrid search wins; the English re-ranker degrades results and was removed. The benchmark "
     "caught this — intuition would have kept it.")
if (FIGS / "fig_methods.png").exists():
    doc.add_picture(str(FIGS / "fig_methods.png"), width=Inches(4.8))

H("10.3 Answer quality, efficiency, and the citation fix", 2)
bullet("the final memoranda were rated ~4.6 / 5 by the LLM judge (legal quality).", "Answer quality: ")
bullet("~2–3 minutes and ~$0.12 per full answer — practical for research use.", "Efficiency: ")
bullet("tightening citations (cite only directly-applicable, verified articles) raised citation-F1 "
       "roughly 3× (0.04 → 0.13) on the same questions.", "Citation precision fix (measured): ")

# ── 11. Findings ───────────────────────────────────────────────────────────────
H("11. Key Findings", 1)
bullet("specialised agents produce grounded, well-structured answers (~4.6/5).", "Multi-agent works: ")
bullet("hybrid beats semantic; the re-ranker hurts; the embedding is validated — all decided by the "
       "benchmark.", "Evidence-based engineering: ")
bullet("answer quality is bounded by retrieval recall — if the correct article is not retrieved, it "
       "cannot be cited. Retrieval is therefore the top priority.", "Bottleneck identified: ")

# ── 12. Engineering ────────────────────────────────────────────────────────────
H("12. Engineering & Reproducibility", 1)
bullet("standardised on Claude Sonnet 5 (temperature 0, retrieval score threshold 0.7 by default); "
       "selectable per run — Sonnet 4.6 / 4.5, Opus 4.6, Haiku 4.5. Reasoning models that reject a "
       "temperature parameter and return content as blocks are handled transparently.", "Model: ")
bullet("the full pipeline also runs headless (without the UI) for batch evaluation.", "Headless: ")
bullet("unit tests and GitHub Actions run automatically on every change (continuous integration).", "Tests + CI: ")
bullet("tokens, cost, and latency are tracked per agent.", "Telemetry: ")
bullet("one command rebuilds the search index; the benchmark can be regenerated on demand.", "Reproducible: ")
bullet("all work is version-controlled and pushed to GitHub.", "Version control: ")

# ── 13. Limitations & future ───────────────────────────────────────────────────
H("13. Limitations & Future Work", 1)
bullet("English/French retrieval lags Arabic (~40% vs. 72%); the corpus covers criminal law only "
       "(Penal Code + Code of Criminal Procedure), and the Criminal Procedure Code is Arabic-only for "
       "now (answers are still given in the user's language).", "Limitations: ")
bullet("evaluate the built-in cross-lingual translation and test stronger multilingual embeddings.",
       "Next — retrieval: ")
bullet("add contract law + French texts (the pipeline supports it); collect expert reference answers at scale.",
       "Next — corpus & rigour: ")
bullet("structured reasoning, a verifier agent, and a full statistical multi-agent-vs-baseline run.",
       "Next — agents & evaluation: ")

# ── 14. Conclusion ─────────────────────────────────────────────────────────────
H("14. Conclusion & Contributions", 1)
bullet("a working trilingual legal AI for Lebanese law: grounded, cited, and adaptive to the user.")
bullet("a rigorous benchmark of grounded questions with validated gold answers and multi-metric scoring.")
bullet("reference-based evaluation — answers judged against a human source of truth, not the judge's opinion.")
bullet("evidence-based design — every retrieval choice proven or rejected by measurement.")
bullet("trust by construction — grounding, citation verification, and a reported hallucination rate.")
bullet("reproducible and engineered — tests, CI, telemetry, and a usable web application.")

# ── Appendix ───────────────────────────────────────────────────────────────────
H("Appendix — Running the Project", 1)
para("Web app:", bold=True, space=2)
code_block([
    "cd code",
    "source venv/bin/activate",
    "streamlit run app.py            # http://localhost:8501",
])
para("First-time setup (fresh machine):", bold=True, space=2)
code_block([
    "python -m venv venv && source venv/bin/activate",
    "pip install -r requirements.txt",
    "cp .env.example .env            # set ANTHROPIC_API_KEY",
    "python scripts/build_index.py   # build the search index",
])
para("Key scripts:", bold=True, space=2)
bullet("build_index.py — build the search index from the corpus.")
bullet("generate_benchmark.py — generate a grounded benchmark dataset.")
bullet("eval_retrieval.py — retrieval metrics (no API cost).")
bullet("run_study.py — multi-agent vs. baselines with the judge.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Saved: {OUT.resolve()}")
print(f"Paragraphs: {len(doc.paragraphs)} | Tables: {len(doc.tables)}")
